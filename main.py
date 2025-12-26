import streamlit as st
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment
from openpyxl.utils import range_boundaries
from io import BytesIO
import os
import base64
from datetime import datetime
import json
from openai import OpenAI
from dotenv import load_dotenv
from pathlib import Path
from typing import List
from docx import Document

# =========================================================
# PATHS robustos (local + deploy)
# =========================================================
BASE_DIR = Path(__file__).resolve().parent

# Carrega .env da mesma pasta do script, independentemente de onde você executar
load_dotenv(dotenv_path=BASE_DIR / ".env")

# Pasta de modelos (permite override no deploy sem mexer no código)
# Exemplo no deploy:
#   MODELOS_DIR="contratos_cadastro"
# ou caminho absoluto:
#   MODELOS_DIR="/mount/src/app/contratos_cadastro"
MODELOS_DIR = Path(os.getenv("MODELOS_DIR", str(BASE_DIR / "contratos_cadastro"))).expanduser().resolve()

def modelo_path(nome_arquivo: str) -> str:
    """Resolve caminho absoluto do arquivo de modelo dentro de MODELOS_DIR."""
    return str((MODELOS_DIR / nome_arquivo).resolve())

def asset_path(*parts) -> str:
    """Resolve caminho absoluto de um asset relativo ao BASE_DIR."""
    return str((BASE_DIR.joinpath(*parts)).resolve())

# Caminho da logomarca (opcional)
logo_path = asset_path("MP.png")


def get_image_base64(file_path: str) -> str:
    if not os.path.exists(file_path):
        return ""
    with open(file_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()


def formatar_data_extenso(data_str: str, cidade: str, uf: str) -> str:
    try:
        meses_pt = {
            1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril", 5: "maio", 6: "junho",
            7: "julho", 8: "agosto", 9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
        }
        data = datetime.strptime(data_str, "%d/%m/%Y")
        dia = data.day
        mes = meses_pt[data.month]
        ano = data.year
        return f"{cidade} - {uf}, {dia:02d} de {mes} de {ano}."
    except Exception:
        return f"{cidade} - {uf}"


def get_top_left_of_merged_cell(ws, cell_coordinate):
    for merged_range in ws.merged_cells.ranges:
        if cell_coordinate in merged_range:
            min_col, min_row, _, _ = range_boundaries(str(merged_range))
            return ws.cell(row=min_row, column=min_col)
    return ws[cell_coordinate]


image_base64 = get_image_base64(logo_path)
if image_base64:
    st.markdown(
        f"""
        <div style="display: flex; justify-content: center; align-items: center; margin-bottom: 20px;">
            <img src="data:image/png;base64,{image_base64}" alt="Logomarca" style="width: 300px;">
        </div>
        """,
        unsafe_allow_html=True,
    )

# ===========
# Auto-fill: inputs via session_state (permite preencher por IA)
# ===========
def inp(chave: str, label: str, placeholder: str = "") -> str:
    if chave not in st.session_state:
        st.session_state[chave] = ""
    return st.text_input(label, value=st.session_state[chave], placeholder=placeholder, key=chave)


def _merge_session_state(d: dict):
    """Atualiza session_state apenas com valores não vazios."""
    for k, v in d.items():
        if v is None:
            continue
        v_str = str(v).strip()
        if v_str:
            st.session_state[k] = v_str


# ===========
# LLM: extrair dados do cliente a partir de PDF
# ===========
SCHEMA_DADOS_CLIENTE = {
    "name": "dados_cliente",
    "schema": {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "IDENTIFICACAO_CIVIL": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "CLIENTE": {"type": ["string", "null"]},
                    "ESTADO_CIVIL": {"type": ["string", "null"]},
                    "DATA_NASC": {"type": ["string", "null"]},
                    "PROFISSAO": {"type": ["string", "null"]},
                    "RG": {"type": ["string", "null"]},
                    "ORGAO_EXPEDIDOR": {"type": ["string", "null"]},
                    "CPF": {"type": ["string", "null"]},
                },
                "required": ["CLIENTE", "ESTADO_CIVIL", "DATA_NASC", "PROFISSAO", "RG", "ORGAO_EXPEDIDOR", "CPF"],
            },
            "ENDERECO": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "ENDERECO": {"type": ["string", "null"]},
                    "Nº": {"type": ["string", "null"]},
                    "BAIRRO": {"type": ["string", "null"]},
                    "COMPLEMENTO": {"type": ["string", "null"]},
                    "CEP": {"type": ["string", "null"]},
                    "CIDADE": {"type": ["string", "null"]},
                    "UF": {"type": ["string", "null"]},
                },
                "required": ["ENDERECO", "Nº", "BAIRRO", "COMPLEMENTO", "CEP", "CIDADE", "UF"],
            },
            "CONTATO": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "EMAIL": {"type": ["string", "null"]},
                    "WHATSAPP": {"type": ["string", "null"]},
                    "TELEFONE2": {"type": ["string", "null"]},
                    "TELEFONE3": {"type": ["string", "null"]},
                },
                "required": ["EMAIL", "WHATSAPP", "TELEFONE2", "TELEFONE3"],
            },
        },
        "required": ["IDENTIFICACAO_CIVIL", "ENDERECO", "CONTATO"],
    },
    "strict": True,
}


def extrair_dados_cliente_de_pdf(pdf_bytes: bytes) -> dict:
    """
    Envia o PDF para a LLM e retorna um dict com:
    IDENTIFICACAO_CIVIL, ENDERECO, CONTATO.
    """
    import io  # necessário para enviar bytes como arquivo

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY não encontrada. Configure no .env (na mesma pasta do script) ou no ambiente.")

    client = OpenAI(api_key=api_key)

    # =========================
    # PASSO 1) Upload do PDF e obter file_id
    # =========================
    uploaded = client.files.create(
        file=("documentos.pdf", io.BytesIO(pdf_bytes)),
        purpose="assistants"
    )
    file_id = uploaded.id

    prompt = """
Você receberá um PDF contendo documentos (ex.: RG e comprovante de residência).
Extraia APENAS os dados do CLIENTE (pessoa atendida), nos grupos:

1) IDENTIFICAÇÃO CIVIL: CLIENTE, ESTADO_CIVIL, DATA_NASC, PROFISSAO, RG, ORGAO_EXPEDIDOR, CPF
2) ENDEREÇO: ENDERECO, Nº, BAIRRO, COMPLEMENTO, CEP, CIDADE, UF
3) CONTATO: EMAIL, WHATSAPP, TELEFONE2, TELEFONE3

REGRAS:
- Não invente nada. Se não encontrar, retorne null.
- Não inclua declarante, testemunhas, rogatário ou terceiros.
- Se houver divergência entre documentos, priorize o mais recente e/ou o que estiver mais legível.
"""

    # =========================
    # PASSO 2) Referenciar o PDF por file_id (sem file_data/base64)
    # =========================
    resp = client.responses.create(
        model="gpt-4o-mini",
        input=[{
            "role": "user",
            "content": [
                {"type": "input_file", "file_id": file_id},
                {"type": "input_text", "text": prompt},
            ]
        }],
        text={
            "format": {
                "type": "json_schema",
                "name": SCHEMA_DADOS_CLIENTE["name"],
                "schema": SCHEMA_DADOS_CLIENTE["schema"],
                "strict": SCHEMA_DADOS_CLIENTE.get("strict", True),
            }
        }
    )

    raw = resp.output_text or ""
    try:
        data = json.loads(raw)
    except Exception:
        # fallback: tenta extrair um json mesmo se vier “sujo”
        ini = raw.find("{")
        fim = raw.rfind("}")
        if ini >= 0 and fim > ini:
            data = json.loads(raw[ini:fim + 1])
        else:
            raise ValueError("Não foi possível interpretar o JSON retornado pela LLM.")

    return data


def importar_pdf_e_preencher(pdf_bytes: bytes):
    data = extrair_dados_cliente_de_pdf(pdf_bytes)

    merged = {}
    merged.update(data.get("IDENTIFICACAO_CIVIL", {}))
    merged.update(data.get("ENDERECO", {}))
    merged.update(data.get("CONTATO", {}))

    _merge_session_state(merged)


st.title("Gerador de Documentos Jurídicos com I.A.")

st.subheader("📎 Importar dados do cliente (PDF)")

pdf_docs = st.file_uploader(
    "Envie um PDF com RG/CNH e comprovante de residência (pode conter mais de 1 página).",
    type=["pdf"]
)

col_imp1, col_imp2 = st.columns([1, 2])
with col_imp1:
    if st.button("🤖 Extrair e preencher", use_container_width=True) and pdf_docs is not None:
        try:
            importar_pdf_e_preencher(pdf_docs.getvalue())
            st.success("✅ Dados importados! Conferindo os campos...")
            st.rerun()
        except Exception as e:
            st.error(f"❌ Falha ao importar do PDF: {e}")

with col_imp2:
    st.caption("Dica: após importar, revise os campos antes de gerar o documento.")


st.subheader("Preencha os dados do cliente")


def carregar_variaveis():
    dados = {
        "CLIENTE": inp("CLIENTE", "Nome do Cliente"),
        "ESTADO_CIVIL": inp("ESTADO_CIVIL", "Estado Civil"),
        "DATA_NASC": inp("DATA_NASC", "Data de Nascimento"),
        "PROFISSAO": inp("PROFISSAO", "Profissão"),
        "RG": inp("RG", "RG"),
        "ORGAO_EXPEDIDOR": inp("ORGAO_EXPEDIDOR", "Órgão Expedidor"),
        "CPF": inp("CPF", "CPF"),

        "ENDERECO": inp("ENDERECO", "Endereço"),
        "Nº": inp("Nº", "Número"),
        "BAIRRO": inp("BAIRRO", "Bairro"),
        "COMPLEMENTO": inp("COMPLEMENTO", "Complemento"),
        "CEP": inp("CEP", "CEP"),
    }

    # Cidade/Comarca com a lógica atual (checkbox), mas agora preenchíveis
    col1, col2 = st.columns([1, 1])
    with col2:
        cidade_editavel = st.checkbox("Editar cidade manualmente?", key="editar_cidade")
    with col1:
        if "CIDADE" not in st.session_state:
            st.session_state["CIDADE"] = "Manaus"
        dados["CIDADE"] = st.text_input(
            "Cidade",
            value=st.session_state["CIDADE"] if not cidade_editavel else "",
            key="CIDADE"
        )

    col3, col4 = st.columns([1, 1])
    with col4:
        comarca_editavel = st.checkbox("Editar comarca manualmente?", key="editar_comarca")
    with col3:
        if "COMARCA" not in st.session_state:
            st.session_state["COMARCA"] = st.session_state.get("CIDADE", "Manaus")
        dados["COMARCA"] = st.text_input(
            "Comarca",
            value=st.session_state["COMARCA"] if not comarca_editavel else "",
            key="COMARCA"
        )

    # ✅ RESTAURADO: campos administrativos + declarante + testemunhas
    dados.update({
        "UF": inp("UF", "UF"),
        "DATA": inp("DATA", "Data"),
        "EMAIL": inp("EMAIL", "Email"),
        "WHATSAPP": inp("WHATSAPP", "WhatsApp"),
        "TELEFONE2": inp("TELEFONE2", "Telefone 2"),
        "TELEFONE3": inp("TELEFONE3", "Telefone 3"),
        "SENHA_GOV": inp("SENHA_GOV", "Senha GOV"),

        "INDICACAO_CLIENTE": inp("INDICACAO_CLIENTE", "Indicação do Cliente"),
        "PARCERIA_ADVOGADO": inp("PARCERIA_ADVOGADO", "Parceria Advogado"),
        "ATENDENTE": inp("ATENDENTE", "Atendente"),

        "DECLARANTE": inp("DECLARANTE", "Nome do Declarante"),
        "DECLARANTE_ESTADO_CIVIL": inp("DECLARANTE_ESTADO_CIVIL", "Estado Civil do Declarante"),
        "DECLARANTE_PROFISSAO": inp("DECLARANTE_PROFISSAO", "Profissão do Declarante"),

        "TESTEMUNHA1": inp("TESTEMUNHA1", "TESTEMUNHA 1 - Nome"),
        "TESTEMUNHA1_CPF": inp("TESTEMUNHA1_CPF", "TESTEMUNHA 1 - CPF"),
        "TESTEMUNHA1_RG": inp("TESTEMUNHA1_RG", "TESTEMUNHA 1 - RG"),
        "TESTEMUNHA1_END": inp("TESTEMUNHA1_END", "TESTEMUNHA 1 - Endereço"),

        "TESTEMUNHA2": inp("TESTEMUNHA2", "TESTEMUNHA 2 - Nome"),
        "TESTEMUNHA2_CPF": inp("TESTEMUNHA2_CPF", "TESTEMUNHA 2 - CPF"),
        "TESTEMUNHA2_RG": inp("TESTEMUNHA2_RG", "TESTEMUNHA 2 - RG"),
        "TESTEMUNHA2_END": inp("TESTEMUNHA2_END", "TESTEMUNHA 2 - Endereço"),
    })

    return dados


dados = carregar_variaveis()

# Pergunta: Cliente alfabetizado?
st.subheader("Cliente alfabetizado(a)?")

opcao_alfabetizado = st.radio(
    "Selecione uma opção:",
    ["Sim", "Não"],
    index=0,
    horizontal=True
)

# Define os modelos de acordo com a escolha
if opcao_alfabetizado == "Sim":
    modelos_arquivo = {
        "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS": modelo_path("CONTR.PREST.SERV.ADV.xlsx"),
        "DECLARAÇÃO DE HIPOSSUFICIÊNCIA DE RENDA": modelo_path("DECLARAÇÃO DE HIPOSSUFICIÊNCIA DE RENDA.xlsx"),
        "DECLARAÇÃO DE VIDA E RESIDÊNCIA": modelo_path("DECLAR.VIDA.RESIDÊNCIA.xlsx"),
        "PROCURAÇÃO AD JUDICIA ET EXTRA": modelo_path("PROCURAÇÃO AD JUDICIA ET EXTRA.xlsx")
    }
else:
    modelos_arquivo = {
        "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS_ANALFABETO(A)": modelo_path("CONTR.PREST.SERV.ADV_ANALFABETO(A).xlsx"),
        "DECLARAÇÃO DE HIPOSSUFICIÊNCIA DE RENDA_ANALFABETO(A)": modelo_path("DECLARAÇÃO DE HIPOSSUFICIÊNCIA DE RENDA_ANALFABETO(A).xlsx"),
        "DECLARAÇÃO DE VIDA E RESIDÊNCIA_ANALFABETO(A)": modelo_path("DECLARAÇÃO_DE_VIDA_E_RESIDÊNCIA_ANALFABETO(A).xlsx"),
        "PROCURAÇÃO AD JUDICIA ET EXTRA_ANALFABETO(A)": modelo_path("PROCURAÇÃO AD JUDICIA ET EXTRA_ANALFABETO(A).xlsx")
    }

    # Campos adicionais se o cliente NÃO for alfabetizado
    dados_rogratario = {}
    if opcao_alfabetizado == "Não":
        st.subheader("Preencha os dados do Rogatário (representante do declarante analfabeto)")
        dados_rogratario["ROGATARIO_NOME"] = inp("ROGATARIO_NOME", "Nome do Rogatário")
        dados_rogratario["ROGATARIO_RG"] = inp("ROGATARIO_RG", "RG do Rogatário")
        dados_rogratario["ROGATARIO_CPF"] = inp("ROGATARIO_CPF", "CPF do Rogatário")
        dados_rogratario["ROGATARIO_END"] = inp("ROGATARIO_END", "Endereço do Rogatário")

        # Atualiza no dicionário principal
        dados.update(dados_rogratario)

modelo_escolhido = st.selectbox("Modelo disponível", list(modelos_arquivo.keys()))

if st.button("Gerar documento preenchido"):
    caminho_excel = modelos_arquivo.get(modelo_escolhido)

    if caminho_excel and os.path.exists(caminho_excel):
        wb = load_workbook(caminho_excel)
        ws = wb.active

        fonte_padrao = Font(name='Arial', size=11)
        fonte_sublinhada = Font(name='Arial', size=11, underline='single')
        alinhamento_esquerda = Alignment(horizontal='left')
        alinhamento_justificado_topo = Alignment(horizontal='justify', vertical='top', wrap_text=True)
        alinhamento_direita_topo = Alignment(horizontal='right', vertical='top')

        if modelo_escolhido == "DECLARAÇÃO DE HIPOSSUFICIÊNCIA DE RENDA":
            texto = (
                f"Eu, {dados['CLIENTE']}, brasileiro (a), estado civil {dados['ESTADO_CIVIL']}, Profissão {dados['PROFISSAO']}, "
                f"RG {dados['RG']} SSP {dados['UF']}. Inscrito no CPF sob o nº {dados['CPF']} residente e domiciliado nesta cidade de "
                f"{dados['CIDADE']}/{dados['UF']} na {dados['ENDERECO']} - {dados['COMPLEMENTO']} Nº {dados['Nº']}, Bairro: {dados['BAIRRO']}, CEP:{dados['CEP']}"
            )
            ws["A8"] = texto
            ws["A8"].font = fonte_padrao
            ws["A8"].alignment = alinhamento_justificado_topo
            ws.row_dimensions[8].height = 30

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["D25"] = data_ext
            ws["D25"].font = fonte_sublinhada
            ws["D25"].alignment = alinhamento_direita_topo

        elif modelo_escolhido == "DECLARAÇÃO DE VIDA E RESIDÊNCIA":
            texto = (
                f"Eu, {dados['DECLARANTE']}. Brasileiro (a), estado civil {dados['DECLARANTE_ESTADO_CIVIL']}, Profissão {dados['DECLARANTE_PROFISSAO']}.\n"
                f"DECLARO para os devidos fins de comprovação de residência, sob as penas da lei (art. 2º da lei 7.115/83), que {dados['CLIENTE']}, "
                f"Brasileiro (a), Estado Civil {dados['ESTADO_CIVIL']}, Profissão: {dados['PROFISSAO']}, portador(a), do RG: {dados['RG']} e CPF: {dados['CPF']}, "
                f"é residente e domiciliado na {dados['ENDERECO']} - Nº {dados['Nº']} - {dados['COMPLEMENTO']}. Bairro: {dados['BAIRRO']}. "
                f"CEP: {dados['CEP']}. Cidade: {dados['CIDADE']}, UF: {dados['UF']}."
            )
            ws["A8"] = texto
            ws["A8"].font = fonte_padrao
            ws["A8"].alignment = alinhamento_justificado_topo
            ws.row_dimensions[8].height = 60

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["I17"] = data_ext
            ws["I17"].font = fonte_padrao
            ws["I17"].alignment = alinhamento_direita_topo

        elif modelo_escolhido == "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS":
            preenchimentos = {
                "C8": "CLIENTE",
                "B10": "ESTADO_CIVIL",
                "G10": "PROFISSAO",
                "B12": "RG",
                "F12": "ORGAO_EXPEDIDOR",
                "H12": "CPF",
                "B14": "ENDERECO",
                "B16": "CEP",
                "H14": "COMPLEMENTO",
                "F1": "INDICACAO_CLIENTE",
                "F2": "PARCERIA_ADVOGADO",
                "F3": "ATENDENTE",
                "E16": "CIDADE",
                "I16": "UF",
                "B18": "EMAIL",
                "B20": "WHATSAPP",
                "E20": "TELEFONE2",
                "H20": "TELEFONE3",
                "G68": "CIDADE",
                "B21": "SENHA_GOV",
                "B84": "TESTEMUNHA1",
                "B85": "TESTEMUNHA1_CPF",
                "B86": "TESTEMUNHA1_RG",
                "B87": "TESTEMUNHA1_END",
                "G84": "TESTEMUNHA2",
                "G85": "TESTEMUNHA2_CPF",
                "G86": "TESTEMUNHA2_RG",
                "G87": "TESTEMUNHA2_END"
            }

            for celula, campo in preenchimentos.items():
                valor = dados.get(campo, "")
                target_cell = get_top_left_of_merged_cell(ws, celula)
                target_cell.value = valor
                target_cell.font = fonte_padrao
                target_cell.alignment = alinhamento_esquerda

            clausula = (
                "I) 30% (trinta por cento) sobre o valor recebido pelo CONTRATANTE em razão de SENTENÇA ou ACORDO, seja este JUDICIAL ou EXTRAJUDICIAL;"
                if dados.get("CIDADE", "").lower() == "manaus" else
                "I) 35% (trinta e cinco por cento) sobre o valor recebido pelo CONTRATANTE em razão de SENTENÇA ou ACORDO, seja este JUDICIAL ou EXTRAJUDICIAL;"
            )
            ws["A29"] = clausula
            ws["A29"].font = fonte_padrao
            ws["A29"].alignment = alinhamento_justificado_topo

            for linha in ["A31", "A34", "A61", "A63", "A65", "B84", "B85", "B86", "B87", "G84", "G85", "G86", "G87"]:
                ws[linha].font = fonte_padrao
                ws[linha].alignment = alinhamento_justificado_topo

            clausula_comarca = f"CLÁUSULA 17ª - As  partes  contratantes  elegem  o  foro  da  Comarca  de {dados.get('COMARCA', '')} para dirimir quaisquer controvérsias oriundas do presente contrato."
            ws["A68"] = clausula_comarca
            ws["A68"].font = fonte_padrao
            ws["A68"].alignment = alinhamento_justificado_topo

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["F72"] = data_ext
            ws["F72"].font = fonte_padrao
            ws["F72"].alignment = alinhamento_direita_topo

            ws["A76"] = dados.get("CLIENTE", "")
            ws["A76"].font = fonte_padrao
            ws["A76"].alignment = alinhamento_justificado_topo

        elif modelo_escolhido == "PROCURAÇÃO AD JUDICIA ET EXTRA":
            preenchimentos = {
                "C3": "CLIENTE",
                "C4": "ESTADO_CIVIL",
                "G4": "PROFISSAO",
                "B5": "RG",
                "G5": "ORGAO_EXPEDIDOR",
                "B6": "CPF",
                "B9": "EMAIL"
            }

            for celula, campo in preenchimentos.items():
                valor = dados.get(campo, "")
                target_cell = get_top_left_of_merged_cell(ws, celula)
                target_cell.value = valor
                target_cell.font = fonte_padrao
                target_cell.alignment = alinhamento_esquerda

            # Construir endereço completo manualmente e aplicar na célula A8
            endereco_formatado = (
                f"{dados.get('ENDERECO', '')}, Nº {dados.get('Nº', '')}, "
                f"{dados.get('BAIRRO', '')}, CEP: {dados.get('CEP', '')}, {dados.get('COMPLEMENTO', '')}"
            )
            ws["A8"] = endereco_formatado
            ws["A8"].font = fonte_padrao
            ws["A8"].alignment = alinhamento_justificado_topo

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["I17"] = data_ext
            ws["I17"].font = fonte_padrao
            ws["I17"].alignment = alinhamento_direita_topo

        elif modelo_escolhido == "DECLARAÇÃO DE VIDA E RESIDÊNCIA_ANALFABETO(A)":
            texto = (
                f"Eu, {dados['DECLARANTE']}. Brasileiro (a), estado civil {dados['DECLARANTE_ESTADO_CIVIL']}, Profissão {dados['DECLARANTE_PROFISSAO']}.\n"
                f"DECLARO para os devidos fins de comprovação de residência, sob as penas da lei (art. 2º da lei 7.115/83), que {dados['CLIENTE']}, "
                f"Brasileiro (a), Estado Civil {dados['ESTADO_CIVIL']}, Profissão: {dados['PROFISSAO']}, portador(a), do RG: {dados['RG']} e CPF: {dados['CPF']}, "
                f"é residente e domiciliado na {dados['ENDERECO']} - Nº {dados['Nº']} - {dados['COMPLEMENTO']}. Bairro: {dados['BAIRRO']}. "
                f"CEP: {dados['CEP']}. Cidade: {dados['CIDADE']}, UF: {dados['UF']}."
            )
            ws["A6"] = texto
            ws["A6"].font = fonte_padrao
            ws["A6"].alignment = alinhamento_justificado_topo
            ws.row_dimensions[8].height = 60

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["I18"] = data_ext
            ws["I18"].font = fonte_padrao
            ws["I18"].alignment = alinhamento_direita_topo

            preenchimentos = {
                "B26": "ROGATARIO_NOME",
                "B27": "ROGATARIO_RG",
                "B28": "ROGATARIO_CPF",
                "B29": "ROGATARIO_END",
                "B32": "TESTEMUNHA1",
                "B34": "TESTEMUNHA1_CPF",
                "B33": "TESTEMUNHA1_RG",
                "B35": "TESTEMUNHA1_END",
                "B38": "TESTEMUNHA2",
                "B40": "TESTEMUNHA2_CPF",
                "B39": "TESTEMUNHA2_RG",
                "B41": "TESTEMUNHA2_END"
            }

            for celula, campo in preenchimentos.items():
                valor = dados.get(campo, "")
                target_cell = get_top_left_of_merged_cell(ws, celula)
                target_cell.value = valor
                target_cell.font = fonte_padrao
                target_cell.alignment = alinhamento_esquerda

        elif modelo_escolhido == "PROCURAÇÃO AD JUDICIA ET EXTRA_ANALFABETO(A)":
            preenchimentos = {
                "C3": "CLIENTE",
                "C4": "ESTADO_CIVIL",
                "F4": "PROFISSAO",
                "B5": "RG",
                "G5": "ORGAO_EXPEDIDOR",
                "B6": "CPF",
                "B9": "EMAIL"
            }

            for celula, campo in preenchimentos.items():
                valor = dados.get(campo, "")
                target_cell = get_top_left_of_merged_cell(ws, celula)
                target_cell.value = valor
                target_cell.font = fonte_padrao
                target_cell.alignment = alinhamento_esquerda

            endereco_formatado = (
                f"{dados.get('ENDERECO', '')}, Nº {dados.get('Nº', '')}, "
                f"{dados.get('BAIRRO', '')}, CEP: {dados.get('CEP', '')}, {dados.get('COMPLEMENTO', '')}"
            )
            ws["A8"] = endereco_formatado
            ws["A8"].font = fonte_padrao
            ws["A8"].alignment = alinhamento_justificado_topo

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["K17"] = data_ext
            ws["K17"].font = fonte_padrao
            ws["K17"].alignment = alinhamento_direita_topo

            # Preencher a célula E23 com: "NOME, CPF: xxx.xxx.xxx-xx"
            nome_roga = dados.get("ROGATARIO_NOME", "").strip()
            cpf_roga = dados.get("ROGATARIO_CPF", "").strip()
            valor_e23 = f"{nome_roga}, CPF: {cpf_roga}" if nome_roga else ""
            target_cell = get_top_left_of_merged_cell(ws, "E23")
            target_cell.value = valor_e23
            target_cell.font = fonte_padrao
            target_cell.alignment = alinhamento_esquerda

            # Preencher as demais células
            preenchimentos = {
                "B26": "TESTEMUNHA1",
                "B28": "TESTEMUNHA1_CPF",
                "B27": "TESTEMUNHA1_END",
                "G26": "TESTEMUNHA2",
                "G28": "TESTEMUNHA2_CPF",
                "G27": "TESTEMUNHA2_END"
            }

            for celula, campo in preenchimentos.items():
                valor = dados.get(campo, "")
                target_cell = get_top_left_of_merged_cell(ws, celula)
                target_cell.value = valor
                target_cell.font = fonte_padrao
                target_cell.alignment = alinhamento_esquerda

        elif modelo_escolhido == "DECLARAÇÃO DE HIPOSSUFICIÊNCIA DE RENDA_ANALFABETO(A)":
            texto = (
                f"Eu, {dados['CLIENTE']}, brasileiro (a), estado civil {dados['ESTADO_CIVIL']}, Profissão {dados['PROFISSAO']}, "
                f"RG {dados['RG']} SSP {dados['UF']}. Inscrito no CPF sob o nº {dados['CPF']} residente e domiciliado nesta cidade de "
                f"{dados['CIDADE']}/{dados['UF']} na {dados['ENDERECO']} - {dados['COMPLEMENTO']} Nº {dados['Nº']}, Bairro: {dados['BAIRRO']}, CEP:{dados['CEP']}"
            )
            ws["A6"] = texto
            ws["A6"].font = fonte_padrao
            ws["A6"].alignment = alinhamento_justificado_topo
            ws.row_dimensions[8].height = 30

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["K20"] = data_ext
            ws["K20"].font = fonte_sublinhada
            ws["K20"].alignment = alinhamento_direita_topo

            preenchimentos = {
                "B24": "ROGATARIO_NOME",
                "B25": "ROGATARIO_RG",
                "B26": "ROGATARIO_CPF",
                "B27": "ROGATARIO_END",
                "B30": "TESTEMUNHA1",
                "B31": "TESTEMUNHA1_RG",
                "B32": "TESTEMUNHA1_CPF",
                "B33": "TESTEMUNHA1_END",
                "B36": "TESTEMUNHA2",
                "B37": "TESTEMUNHA2_RG",
                "B38": "TESTEMUNHA2_CPF",
                "B39": "TESTEMUNHA2_END"
            }

            for celula, campo in preenchimentos.items():
                valor = dados.get(campo, "")
                target_cell = get_top_left_of_merged_cell(ws, celula)
                target_cell.value = valor
                target_cell.font = fonte_padrao
                target_cell.alignment = alinhamento_esquerda

        elif modelo_escolhido == "CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS_ANALFABETO(A)":
            preenchimentos = {
                "C8": "CLIENTE",
                "B10": "ESTADO_CIVIL",
                "G10": "PROFISSAO",
                "B12": "RG",
                "F12": "ORGAO_EXPEDIDOR",
                "H12": "CPF",
                "B14": "ENDERECO",
                "B16": "CEP",
                "G14": "COMPLEMENTO",
                "F1": "INDICACAO_CLIENTE",
                "F2": "PARCERIA_ADVOGADO",
                "F3": "ATENDENTE",
                "E16": "CIDADE",
                "I16": "UF",
                "B18": "EMAIL",
                "B20": "WHATSAPP",
                "E20": "TELEFONE2",
                "H20": "TELEFONE3",
                "G68": "CIDADE",
                "B21": "SENHA_GOV",
                "B79": "ROGATARIO_NOME",
                "B80": "ROGATARIO_RG",
                "B81": "ROGATARIO_CPF",
                "B82": "ROGATARIO_END",
                "B85": "TESTEMUNHA1",
                "B87": "TESTEMUNHA1_CPF",
                "B86": "TESTEMUNHA1_RG",
                "B88": "TESTEMUNHA1_END",
                "B91": "TESTEMUNHA2",
                "B93": "TESTEMUNHA2_CPF",
                "B92": "TESTEMUNHA2_RG",
                "B94": "TESTEMUNHA2_END",
                "B102": "TESTEMUNHA1",
                "B103": "TESTEMUNHA1_CPF",
                "B104": "TESTEMUNHA1_RG",
                "B105": "TESTEMUNHA1_END",
                "G102": "TESTEMUNHA2",
                "G103": "TESTEMUNHA2_CPF",
                "G104": "TESTEMUNHA2_RG",
                "G105": "TESTEMUNHA2_END"
            }

            for celula, campo in preenchimentos.items():
                valor = dados.get(campo, "")
                target_cell = get_top_left_of_merged_cell(ws, celula)
                target_cell.value = valor
                target_cell.font = fonte_padrao
                target_cell.alignment = alinhamento_esquerda

            clausula = (
                "I) 30% (trinta por cento) sobre o valor recebido pelo CONTRATANTE em razão de SENTENÇA ou ACORDO, seja este JUDICIAL ou EXTRAJUDICIAL;"
                if dados.get("CIDADE", "").lower() == "manaus" else
                "I) 35% (trinta e cinco por cento) sobre o valor recebido pelo CONTRATANTE em razão de SENTENÇA ou ACORDO, seja este JUDICIAL ou EXTRAJUDICIAL;"
            )
            ws["A29"] = clausula
            ws["A29"].font = fonte_padrao
            ws["A29"].alignment = alinhamento_justificado_topo

            for linha in ["A31", "A34", "A58", "A64"]:
                ws[linha].font = fonte_padrao
                ws[linha].alignment = alinhamento_justificado_topo

            clausula_comarca = f"CLÁUSULA 17ª - As  partes  contratantes  elegem  o  foro  da  Comarca  de {dados.get('COMARCA', '')} para dirimir quaisquer controvérsias oriundas do presente contrato."
            ws["A67"] = clausula_comarca
            ws["A67"].font = fonte_padrao
            ws["A67"].alignment = alinhamento_justificado_topo

            data_ext = formatar_data_extenso(dados.get("DATA", ""), dados.get("CIDADE", ""), dados.get("UF", ""))
            ws["I71"] = data_ext
            ws["I71"].font = fonte_padrao
            ws["I71"].alignment = alinhamento_direita_topo

            ws["A74"] = dados.get("CLIENTE", "")
            ws["A74"].font = fonte_padrao
            ws["A74"].alignment = alinhamento_justificado_topo

        output = BytesIO()
        wb.save(output)
        output.seek(0)
        st.download_button(
            label="📥 Baixar documento preenchido",
            data=output,
            file_name=f"{modelo_escolhido.replace(' ', '_')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        st.success("✅ Planilha preenchida com sucesso!")
    else:
        st.error("❌ Arquivo de modelo não encontrado.")
        st.caption(f"Procurado em: {caminho_excel}")
        st.caption(f"MODELOS_DIR: {MODELOS_DIR}")

# ==============================
# RECIBOS DE SERVIÇOS JURÍDICOS
# ==============================

# ---------- Utilidades p/ moeda PT-BR ----------
def parse_valor_brl(s: str) -> float:
    """Converte '1.234,56' | '1234,56' | '1234.56' -> 1234.56 (float)."""
    if not s:
        return 0.0
    s = s.strip().replace("R$", "").replace(" ", "")
    s = s.replace(".", "").replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return 0.0


def formatar_brl(v: float) -> str:
    """Formata 1234.56 -> '1.234,56'."""
    return f"{v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _extenso_0_999(n: int) -> str:
    unidades = ["", "um", "dois", "três", "quatro", "cinco", "seis", "sete", "oito", "nove"]
    dez_a_dezenove = ["dez", "onze", "doze", "treze", "quatorze", "quinze", "dezesseis", "dezessete", "dezoito", "dezenove"]
    dezenas = ["", "", "vinte", "trinta", "quarenta", "cinquenta", "sessenta", "setenta", "oitenta", "noventa"]
    centenas = ["", "cento", "duzentos", "trezentos", "quatrocentos", "quinhentos", "seiscentos", "setecentos", "oitocentos", "novecentos"]
    if n == 0:
        return ""
    if n == 100:
        return "cem"
    c = n // 100
    d = (n % 100) // 10
    u = n % 10
    partes = []
    if c:
        partes.append(centenas[c])
    if d == 1:
        partes.append(dez_a_dezenove[u])
    else:
        if d:
            partes.append(dezenas[d])
        if u:
            partes.append(unidades[u])
    saida = ""
    for p in partes:
        saida = p if not saida else f"{saida} e {p}"
    return saida


def _bloco_extenso(n: int, singular: str, plural: str) -> str:
    if n == 0:
        return ""
    if n == 1:
        return f"um {singular}"
    return f"{_extenso_0_999(n)} {plural}"


def numero_para_moeda_ptbr(valor: float) -> str:
    """
    1234.56 -> 'Mil duzentos e trinta e quatro reais e cinquenta e seis centavos'
    (primeira letra maiúscula).
    """
    if valor < 0:
        frase = "menos " + numero_para_moeda_ptbr(-valor)
        return frase[0].upper() + frase[1:]

    inteiro = int(valor)
    centavos = int(round((valor - inteiro) * 100))
    bilhoes = inteiro // 1_000_000_000
    resto = inteiro % 1_000_000_000
    milhoes = resto // 1_000_000
    resto %= 1_000_000
    milhares = resto // 1_000
    centenas = resto % 1_000

    partes = []
    if bilhoes:
        partes.append(_bloco_extenso(bilhoes, "bilhão", "bilhões"))
    if milhoes:
        partes.append(_bloco_extenso(milhoes, "milhão", "milhões"))
    if milhares:
        partes.append("mil" if milhares == 1 else f"{_extenso_0_999(milhares)} mil")
    if centenas:
        partes.append(_extenso_0_999(centenas))

    partes_reais = "zero" if not partes else " ".join(partes).replace("mil e ", "mil ")
    sufx_reais = "real" if inteiro == 1 else "reais"
    frase = f"{partes_reais} {sufx_reais}"

    if centavos:
        ext_cent = _extenso_0_999(centavos)
        sufx_cent = "centavo" if centavos == 1 else "centavos"
        frase += f" e {ext_cent} {sufx_cent}"

    return frase[0].upper() + frase[1:] if frase else frase
# ---------- fim utilidades moeda ----------

st.divider()

# ==============================
# BLOCO — RECIBO (corrigido)
# ==============================

def numero_para_moeda_ptbr(valor: float) -> str:
    """
    1234.56 -> 'Mil duzentos e trinta e quatro reais e cinquenta e seis centavos'
    (primeira letra maiúscula).
    """
    if valor < 0:
        frase = "menos " + numero_para_moeda_ptbr(-valor)
        return frase[0].upper() + frase[1:]

    inteiro = int(valor)
    centavos = int(round((valor - inteiro) * 100))
    bilhoes = inteiro // 1_000_000_000
    resto = inteiro % 1_000_000_000
    milhoes = resto // 1_000_000
    resto %= 1_000_000
    milhares = resto // 1_000
    centenas = resto % 1_000

    partes = []
    if bilhoes:
        partes.append(_bloco_extenso(bilhoes, "bilhão", "bilhões"))
    if milhoes:
        partes.append(_bloco_extenso(milhoes, "milhão", "milhões"))
    if milhares:
        partes.append("mil" if milhares == 1 else f"{_extenso_0_999(milhares)} mil")
    if centenas:
        partes.append(_extenso_0_999(centenas))

    partes_reais = "zero" if not partes else " ".join(partes).replace("mil e ", "mil ")
    sufx_reais = "real" if inteiro == 1 else "reais"
    frase = f"{partes_reais} {sufx_reais}"

    if centavos:
        ext_cent = _extenso_0_999(centavos)
        sufx_cent = "centavo" if centavos == 1 else "centavos"
        frase += f" e {ext_cent} {sufx_cent}"

    return frase[0].upper() + frase[1:] if frase else frase



def _parse_valor_brl(s: str) -> float:
    """Converte '1.234,56' ou '1234,56' em float 1234.56. Vazio/invalid -> 0.0"""
    if not s:
        return 0.0
    s = str(s).strip()
    # mantém só dígitos e separadores
    s = re.sub(r"[^\d,\.]", "", s)
    if not s:
        return 0.0
    # se tem vírgula, assume vírgula como decimal e remove pontos de milhar
    if "," in s:
        s = s.replace(".", "").replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return 0.0


# --- SessionState (para não "resetar" os campos a cada rerun) ---
_recibo_defaults = {
    "recibo_tipo": "CONSULTORIA JURÍDICA",
    "recibo_valor_str": "",
    "recibo_hora": datetime.now().strftime("%H:%M"),
    "recibo_manual_extenso": False,
    "recibo_valor_extenso_manual": "",
    "recibo_preview": "",
}
for _k, _v in _recibo_defaults.items():
    st.session_state.setdefault(_k, _v)

st.subheader("Recibo de serviços jurídicos")

# 1) Tipo
tipo_recibo = st.selectbox(
    "Selecione o tipo de recibo:",
    options=[
        "CONSULTORIA JURÍDICA",
        "ELABORAÇÃO DE PEÇA PROCESSUAL",
        "AUDIÊNCIA",
        "DILIGÊNCIA",
        "OUTROS",
    ],
    key="recibo_tipo",
)

# 2) Valor e Hora
col1, col2 = st.columns(2)
with col1:
    st.text_input("VALOR (R$) — {VALOR}", placeholder="300,00", key="recibo_valor_str")
with col2:
    st.text_input("HORA — {HORA}", placeholder="10:35", key="recibo_hora")

# 3) Extenso (auto / manual)
st.checkbox("Editar valor por extenso manualmente?", key="recibo_manual_extenso")

valor_float = _parse_valor_brl(st.session_state["recibo_valor_str"])
valor_extenso_auto = numero_para_moeda_ptbr(valor_float)

if st.session_state["recibo_manual_extenso"]:
    st.text_input(
        "VALOR_EXTENSO — ({VALOR_EXTENSO})",
        placeholder=valor_extenso_auto,
        key="recibo_valor_extenso_manual",
    )
    valor_extenso_final = st.session_state["recibo_valor_extenso_manual"].strip() or valor_extenso_auto
else:
    st.text_input(
        "VALOR_EXTENSO — ({VALOR_EXTENSO})",
        value=valor_extenso_auto,
        disabled=True,
        key="recibo_valor_extenso_readonly",
    )
    valor_extenso_final = valor_extenso_auto

# 4) Texto base do recibo
nome_cliente = (dados.get("CLIENTE") or st.session_state.get("CLIENTE") or "CLIENTE").strip()
cpf_cliente = (dados.get("CPF") or st.session_state.get("CPF") or "CPF").strip()
data_recibo = datetime.now().strftime("%d/%m/%Y")
hora_recibo = st.session_state["recibo_hora"].strip() or datetime.now().strftime("%H:%M")

valor_brl_format = f"{valor_float:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

texto_base = (
    f"Recebi de {nome_cliente}, portador(a) do CPF {cpf_cliente}, a importância de R$ {valor_brl_format} "
    f"({valor_extenso_final}), face à {tipo_recibo.lower()} realizada no dia {data_recibo}, "
    f"às {hora_recibo} horas, qual dou plena quitação."
)

# 5) Preview editável (persistente)
# Só define o texto inicial uma única vez, para não sobrescrever o que você edita manualmente
if not st.session_state["recibo_preview"]:
    st.session_state["recibo_preview"] = texto_base

st.text_area(
    "Pré-visualização (editável):",
    height=220,
    key="recibo_preview",
)

colb1, colb2 = st.columns([1, 1])
with colb1:
    if st.button("Regerar texto automaticamente"):
        st.session_state["recibo_preview"] = texto_base
        st.rerun()
with colb2:
    if st.button("Limpar campos do recibo"):
        for _k, _v in _recibo_defaults.items():
            st.session_state[_k] = _v
        st.rerun()

# ---- Utilidades de manipulação do .docx
def replace_in_paragraph(paragraph, mapping: dict):
    novo = preencher_texto(paragraph.text, mapping)
    if novo != paragraph.text:
        paragraph.text = novo  # substitui placeholders; perde estilos de runs


def replace_in_table(table, mapping: dict):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, mapping)


def _inserir_3_linhas_apos_titulo(doc: Document, titulo_ref: str, linhas: List[str]):
    """
    Procura o parágrafo com 'titulo_ref' e insere os itens de 'linhas'
    protocolarmente 3 linhas abaixo. Cria parágrafos vazios se necessário.
    """
    titulo_ref_upper = titulo_ref.upper()
    for i, p in enumerate(doc.paragraphs):
        if titulo_ref_upper in (p.text or "").upper():
            insert_index = i + 3
            while len(doc.paragraphs) <= insert_index:
                doc.add_paragraph("")
            for linha in linhas:
                doc.paragraphs[insert_index].insert_paragraph_before(linha)
                insert_index += 1
            break


def render_docx_from_template(
    template_path: str,
    mapping: dict,
    linhas_consultoria: List[str] | None = None,
    data_extenso_str: str = ""
) -> BytesIO:
    doc = Document(template_path)

    # 1) substitui placeholders existentes
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)
    for t in doc.tables:
        replace_in_table(t, mapping)

    # 2) insere o texto 3 linhas abaixo do título
    if linhas_consultoria:
        _inserir_3_linhas_apos_titulo(doc, "RECIBO DE PAGAMENTO", linhas_consultoria)

    # 3) insere {DATA em extenso} 2 linhas acima de "MARCELA DA SILVA PAULO" (à direita)
    if data_extenso_str:
        alvo_upper = "MARCELA DA SILVA PAULO"
        for i, p in enumerate(doc.paragraphs):
            if alvo_upper in (p.text or "").upper():
                insert_index = max(i - 2, 0)
                novo = doc.paragraphs[insert_index].insert_paragraph_before(data_extenso_str)
                novo.alignment = 2  # right
                break

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ---- Botões de geração/Download
colg1, colg2 = st.columns([1, 1])
with colg1:
    gerar = st.button("🔄 Gerar arquivo (.docx)")
with colg2:
    st.write("")

if gerar:
    if not os.path.exists(TEMPLATE_DOCX):
        st.error(
            "❌ Arquivo base não encontrado.\n"
            f"Verifique:\n- {DOCX_PATH_1}\n- {DOCX_PATH_2}"
        )
    else:
        if selecionada == "CONSULTORIA JURÍDICA":
            texto_final_consultoria = preview_editado or preencher_texto(texto_base_consultoria, placeholders)
        else:
            texto_final_consultoria = preencher_texto(texto_base_consultoria, placeholders)

        linhas_para_inserir = [linha for linha in texto_final_consultoria.split("\n")]

        buffer_docx = render_docx_from_template(
            TEMPLATE_DOCX,
            placeholders,
            linhas_consultoria=linhas_para_inserir,
            data_extenso_str=placeholders.get("{DATA em extenso}", "")
        )

        st.success("✅ Arquivo gerado. Clique para baixar:")
        st.download_button(
            label="📥 Baixar Recibo (.docx)",
            data=buffer_docx,
            file_name="Recibo_de_servicos_juridicos.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
